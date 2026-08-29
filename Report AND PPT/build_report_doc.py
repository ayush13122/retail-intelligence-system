import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

OUTPUT_PATH = r"d:\retail_analytics\Ayush_Raj_CSE343_CSE443_Report.docx"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D5D8DC"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_center_para(doc, text, size=12, bold=False, italic=False, space_after=6, space_before=0, color=RGBColor(0x22, 0x22, 0x22)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def build_professional_report():
    doc = Document()

    # Standard 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    # =========================================================================
    # PAGE 1: COVER PAGE (Fixed Spacing, No Soft Returns)
    # =========================================================================
    add_center_para(doc, "SUMMER TRAINING/INTERNSHIP", size=16, bold=True, space_before=36, space_after=6, color=RGBColor(0x1B, 0x36, 0x5D))
    add_center_para(doc, "PROJECT REPORT", size=16, bold=True, space_after=24, color=RGBColor(0x1B, 0x36, 0x5D))
    add_center_para(doc, "(Term June-July 2025 / Academic Session 2025-2026)", size=11, italic=True, space_after=36)
    
    add_center_para(doc, "ENTERPRISE RETAIL INTELLIGENCE &", size=16, bold=True, space_after=6, color=RGBColor(0x0E, 0x4D, 0x92))
    add_center_para(doc, "CUSTOMER SEGMENTATION SYSTEM", size=16, bold=True, space_after=12, color=RGBColor(0x0E, 0x4D, 0x92))
    
    add_center_para(doc, "(Oracle SQL Relational Warehousing, RFM Feature Engineering,", size=11, italic=True, space_after=4)
    add_center_para(doc, "and Unsupervised K-Means Machine Learning)", size=11, italic=True, space_after=48)
    
    add_center_para(doc, "Submitted by", size=12, space_after=6)
    add_center_para(doc, "AYUSH RAJ", size=14, bold=True, space_after=6)
    add_center_para(doc, "Registration Number: 12315673", size=12, space_after=4)
    add_center_para(doc, "Course Code: CSE343 / CSE443", size=12, space_after=48)
    
    add_center_para(doc, "Under the Guidance of", size=12, space_after=6)
    add_center_para(doc, "(Name of Mentor with Designation)", size=12, italic=True, space_after=12)
    add_center_para(doc, "School of Computer Science and Engineering", size=13, bold=True, space_after=4)
    add_center_para(doc, "Lovely Professional University, Phagwara, Punjab", size=13, bold=True, space_after=0)

    doc.add_page_break()

    # =========================================================================
    # PAGE 2: CERTIFICATE
    # =========================================================================
    add_center_para(doc, "CERTIFICATE", size=16, bold=True, space_before=24, space_after=24, color=RGBColor(0x1B, 0x36, 0x5D))

    p_cert = doc.add_paragraph()
    p_cert.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cert.paragraph_format.line_spacing = 1.4
    p_cert.paragraph_format.space_after = Pt(48)
    p_cert.add_run("This is to certify that the project report entitled ")
    p_cert.add_run('"Enterprise Retail Intelligence & Customer Segmentation System" ').bold = True
    p_cert.add_run("submitted by ")
    p_cert.add_run("Ayush Raj ").bold = True
    p_cert.add_run("(Registration Number: ")
    p_cert.add_run("12315673").bold = True
    p_cert.add_run(") in partial fulfillment of the requirements for the Summer Training/Internship course (")
    p_cert.add_run("CSE343 / CSE443").bold = True
    p_cert.add_run(") under the School of Computer Science and Engineering, Lovely Professional University, Phagwara, Punjab, is an authentic record of work carried out by him under supervision.")

    p_sign1 = doc.add_paragraph("Mentor Name & Designation: __________________________")
    p_sign1.paragraph_format.space_after = Pt(24)
    
    p_sign2 = doc.add_paragraph("Signature of Mentor: __________________________")
    p_sign2.paragraph_format.space_after = Pt(36)
    
    p_sign3 = doc.add_paragraph("Head of School")
    p_sign3.paragraph_format.space_after = Pt(4)
    p_sign4 = doc.add_paragraph("School of Computer Science and Engineering")
    p_sign4.paragraph_format.space_after = Pt(4)
    p_sign5 = doc.add_paragraph("Lovely Professional University, Phagwara")
    
    doc.add_page_break()

    # =========================================================================
    # PAGE 3: ACKNOWLEDGEMENT
    # =========================================================================
    add_center_para(doc, "ACKNOWLEDGEMENT", size=16, bold=True, space_before=20, space_after=20, color=RGBColor(0x1B, 0x36, 0x5D))

    p_ack = doc.add_paragraph()
    p_ack.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ack.paragraph_format.line_spacing = 1.3
    p_ack.add_run("I would like to express my sincere gratitude to Lovely Professional University and the Centre for Professional Enhancement (CPE) for organizing the structured skill training in data science and enterprise engineering, which provided me with the opportunity to strengthen my database engineering, mathematical feature extraction, and machine learning skills.")
    
    p_ack2 = doc.add_paragraph()
    p_ack2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ack2.paragraph_format.line_spacing = 1.3
    p_ack2.add_run("I am deeply thankful to my mentor for their constant guidance, valuable architectural feedback, and technical encouragement throughout the development of the Enterprise Retail Intelligence System. Their mentorship helped me master relational Star Schema design, procedural SQL triggers, and unsupervised clustering validation.")
    
    p_ack3 = doc.add_paragraph()
    p_ack3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ack3.paragraph_format.line_spacing = 1.3
    p_ack3.add_run("I also extend my thanks to the faculty members of the School of Computer Science and Engineering and to my peers who assisted during testing and data verification phases.")
    
    p_ack4 = doc.add_paragraph()
    p_ack4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ack4.paragraph_format.line_spacing = 1.3
    p_ack4.add_run("Finally, I extend my heartfelt thanks to my family and friends for their continuous motivation and support throughout the duration of this training.")

    p_ack_sign1 = doc.add_paragraph("Ayush Raj")
    p_ack_sign1.paragraph_format.space_before = Pt(36)
    p_ack_sign1.runs[0].bold = True
    p_ack_sign2 = doc.add_paragraph("Registration No. 12315673")

    doc.add_page_break()

    # =========================================================================
    # PAGE 4: TABLE OF CONTENTS
    # =========================================================================
    add_center_para(doc, "TABLE OF CONTENTS", size=16, bold=True, space_before=16, space_after=20, color=RGBColor(0x1B, 0x36, 0x5D))

    toc_items = [
        ("Certificate", "ii"),
        ("Acknowledgement", "iii"),
        ("Table of Contents", "iv"),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("    1.1 Company / Institute Profile", "1"),
        ("    1.2 Overview of Training Domain", "1"),
        ("    1.3 Objective of the Project", "2"),
        ("CHAPTER 2: TRAINING OVERVIEW", "3"),
        ("    2.1 Tools & Technologies Used", "3"),
        ("    2.2 Areas Covered During Training", "3"),
        ("    2.3 Daily / Weekly Work Summary", "4"),
        ("CHAPTER 3: PROJECT DETAILS", "5"),
        ("    3.1 Title of the Project", "5"),
        ("    3.2 Problem Definition", "5"),
        ("    3.3 Scope and Objectives", "5"),
        ("    3.4 System Requirements", "6"),
        ("    3.5 System Architecture", "6"),
        ("    3.6 Data Flow & Entity Relationship Architecture", "7"),
        ("CHAPTER 4: IMPLEMENTATION", "8"),
        ("    4.1 Tools Used", "8"),
        ("    4.2 Methodology", "8"),
        ("    4.3 Functional Modules", "9"),
        ("    4.4 Code Snippets", "10"),
        ("    4.5 Execution Outputs & Terminal Verification", "11"),
        ("CHAPTER 5: RESULTS AND DISCUSSION", "12"),
        ("    5.1 Output / Business KPIs & Cluster Cohorts", "12"),
        ("    5.2 Challenges Faced", "13"),
        ("    5.3 Learnings & Technical Insights", "13"),
        ("CHAPTER 6: CONCLUSION", "14"),
        ("    6.1 Summary", "14"),
        ("    6.2 Future Enhancements", "14"),
        ("ANNEXURE A: MASTER BATCH PIPELINE SOURCE CODE", "15"),
        ("ANNEXURE B: PRESENTATION DECK STRUCTURE", "17")
    ]
    for title, page in toc_items:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(3)
        r_t1 = p_t.add_run(title)
        if not title.startswith("    "):
            r_t1.bold = True
            r_t1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        dots = "." * (80 - len(title) * 2)
        p_t.add_run(f" {dots} {page}")

    doc.add_page_break()

    # =========================================================================
    # CHAPTER BUILDERS
    # =========================================================================
    def add_chapter_heading(num_str, title_str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"CHAPTER {num_str}: {title_str.upper()}")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    def add_section_heading(sec_str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(sec_str)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x0E, 0x4D, 0x92)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F4F6F7")
        set_cell_margins(cell, 120, 120, 140, 140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        doc.add_paragraph()

    # --- CHAPTER 1 ---
    add_chapter_heading("1", "Introduction")
    add_section_heading("1.1 COMPANY / INSTITUTE PROFILE")
    doc.add_paragraph(
        "This project was undertaken as part of the Summer Training programme organized by the Centre for Professional Enhancement (CPE), "
        "Lovely Professional University (LPU), Phagwara. LPU is a NAAC A++ accredited premier institution. The Centre for Professional Enhancement "
        "delivers industry-aligned technical skill training courses designed to bridge academic concepts with real-world enterprise software engineering."
    )

    add_section_heading("1.2 OVERVIEW OF TRAINING DOMAIN")
    doc.add_paragraph(
        "The training domain encompasses Enterprise Relational Warehousing, Feature Engineering, Applied Machine Learning, and Business Intelligence (BI). "
        "The curriculum focused on designing normalized Star Schemas in SQL, implementing procedural PL/SQL triggers for calculation consistency, "
        "performing multi-dimensional RFM (Recency, Frequency, Monetary) vector transformations, and training unsupervised K-Means clustering algorithms "
        "evaluated through geometric Silhouette scoring coefficients."
    )

    add_section_heading("1.3 OBJECTIVE OF THE PROJECT")
    doc.add_paragraph("The primary technical objectives of the Enterprise Retail Intelligence & Customer Segmentation System are:")
    for obj in [
        "To architect a normalized relational schema with Fact and Dimension tables (dim_customers, dim_products, fact_transactions) enforcing strict referential integrity.",
        "To compile automated row-level PL/SQL triggers that compute total line spend dynamically during insertion without external application overhead.",
        "To aggregate raw transaction timestamps, distinct order IDs, and line items into standardized 3D customer RFM vector representations.",
        "To train an unsupervised K-Means clustering model (K=3) on StandardScaler normalized distributions to segment shoppers into Champions, Potential Loyal, and Hibernating tiers.",
        "To evaluate cluster separation mathematically using Silhouette coefficients and deploy multi-platform presentation layers across Streamlit web and Tkinter desktop interfaces."
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(obj)

    # --- CHAPTER 2 ---
    add_chapter_heading("2", "Training Overview")
    add_section_heading("2.1 TOOLS & TECHNOLOGIES USED")
    for tool in [
        "Programming Languages: Python 3.10+, ANSI SQL / SQLite Engine",
        "Database Engine: SQLite Relational Database Engine (Oracle DDL/DML standard)",
        "Machine Learning Libraries: Scikit-Learn (KMeans, StandardScaler, silhouette_score)",
        "Data Analytics & Math: Pandas, NumPy",
        "Visualization & GUI Platforms: Streamlit Web Framework, Matplotlib, Tkinter Desktop GUI",
        "Development Tools: Visual Studio Code, Integrated Terminal CLI, Git Version Control"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(tool)

    add_section_heading("2.2 AREAS COVERED DURING TRAINING")
    doc.add_paragraph(
        "The training covered relational database architecture, data sanitization pipelines, vector-accelerated feature aggregation, "
        "distance-metric machine learning models, and real-time visualization frameworks. Special focus was given to mitigating feature scale dominance "
        "using Z-score standardization and separating batch data pipelines from interactive frontend dashboards."
    )

    add_section_heading("2.3 DAILY / WEEKLY WORK SUMMARY")
    table1 = doc.add_table(rows=7, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table1)
    headers = ["Week", "Focus Area", "Key Deliverables & Technical Milestones"]
    for i, h in enumerate(headers):
        cell = table1.cell(0, i)
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)

    wk_data = [
        ("Week 1", "Schema Design & DDL Construction", "Designed Star Schema; compiled DDL scripts for dim_customers, dim_products, and fact_transactions."),
        ("Week 2", "PL/SQL Triggers & Advanced SQL", "Built automated trigger trg_calc_transaction_total; developed multi-table joins and CLV aggregations."),
        ("Week 3", "Data Ingestion & ETL Automation", "Built Python ETL scripts to parse transaction datasets, filter invalid customer records, and export clean CSVs."),
        ("Week 4", "RFM Feature Engineering & Normalization", "Engineered Recency, Frequency, and Monetary feature metrics per customer; implemented StandardScaler."),
        ("Week 5", "Unsupervised K-Means Training", "Implemented K-Means (K=3); evaluated Silhouette score (0.3789); mapped behavioral cluster profiles."),
        ("Week 6", "Multi-Tier Frontend & Report Writing", "Developed Streamlit web app, Tkinter desktop GUI, master batch runner, and final documentation.")
    ]
    for row_idx, data in enumerate(wk_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table1.cell(row_idx, col_idx)
            set_cell_background(cell, "FFFFFF" if row_idx % 2 != 0 else "F8F9F9")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(10)
    doc.add_paragraph()

    # --- CHAPTER 3 ---
    add_chapter_heading("3", "Project Details")
    add_section_heading("3.1 TITLE OF THE PROJECT")
    doc.add_paragraph("Enterprise Retail Intelligence & Customer Segmentation System — an end-to-end relational data warehousing, RFM analytics, and machine learning platform.")

    add_section_heading("3.2 PROBLEM DEFINITION")
    doc.add_paragraph(
        "Modern enterprise retail platforms generate massive transaction logs containing missing records, cancellations, and varied pricing. "
        "Standard transactional tables cannot identify at-risk customers or distinguish high-value shoppers. Without automated calculation triggers, "
        "feature normalization, and machine learning segmentation, businesses cannot tailor retention strategies or identify profitable product categories."
    )

    add_section_heading("3.3 SCOPE AND OBJECTIVES")
    doc.add_paragraph(
        "The project provides an enterprise data pipeline that sanitizes transaction records, computes row-level line amounts at the database level, "
        "extracts RFM behavioral vectors, segments shoppers into three distinct tiers using K-Means, and renders metrics across interactive dashboards."
    )

    add_section_heading("3.4 SYSTEM REQUIREMENTS")
    for req in [
        "Hardware: Intel Core i3 / AMD Ryzen 3 or higher, minimum 4 GB RAM, 50 MB disk space.",
        "Operating System: Windows 10/11, Ubuntu Linux, or macOS.",
        "Software & Tools: Python 3.10 - 3.14 (64-bit), VS Code IDE, SQLite Database Engine."
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(req)

    add_section_heading("3.5 SYSTEM ARCHITECTURE")
    doc.add_paragraph("The application follows a decoupled multi-tier architecture separating database storage, machine learning computation, and user presentation:")
    arch_ascii = (
        "+-------------------------------------------------------------------------+\n"
        "|                         PRESENTATION LAYER                              |\n"
        "|   +---------------------------------+  +----------------------------+   |\n"
        "|   | Streamlit Web App (app.py)      |  | Tkinter GUI (app_gui.py)   |   |\n"
        "|   +---------------------------------+  +----------------------------+   |\n"
        "+------------------------------------+------------------------------------+\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|                  ANALYTICAL & MACHINE LEARNING LAYER                    |\n"
        "|   +-----------------------------------------------------------------+   |\n"
        "|   | 1. RFM Feature Extraction (Recency, Frequency, Monetary)        |   |\n"
        "|   | 2. StandardScaler Normalization Engine                          |   |\n"
        "|   | 3. Unsupervised K-Means Clustering (k=3, Silhouette Evaluation) |   |\n"
        "|   +-----------------------------------------------------------------+   |\n"
        "+------------------------------------+------------------------------------+\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|                 DATABASE & PERSISTENT STORAGE LAYER                     |\n"
        "|   +---------------------------------+  +----------------------------+   |\n"
        "|   | Relational Fact/Dim Tables      |  | CSV Pipeline Exports       |   |\n"
        "|   | (dim_customers, dim_products,   |  | (retail_clean.csv,         |   |\n"
        "|   |  fact_transactions, Triggers)   |  |  customer_segments.csv)    |   |\n"
        "|   +---------------------------------+  +----------------------------+   |\n"
        "+-------------------------------------------------------------------------+"
    )
    add_code_block(arch_ascii)

    add_section_heading("3.6 DATA FLOW & ENTITY RELATIONSHIP ARCHITECTURE")
    doc.add_paragraph("The relational schema utilizes Star Schema normalization to maintain clear boundaries between master entities and transactional facts:")
    er_ascii = (
        "+------------------+             +-----------------------+             +------------------+\n"
        "|  DIM_CUSTOMERS   |             |   FACT_TRANSACTIONS   |             |   DIM_PRODUCTS   |\n"
        "+------------------+             +-----------------------+             +------------------+\n"
        "| * customer_id PK | <---------- | * customer_id FK      | ----------> | * stock_code PK  |\n"
        "|   country        | 1         N | * stock_code FK       | N         1 |   description    |\n"
        "|   created_at     |             |   invoice_no          |             |   unit_price     |\n"
        "+------------------+             |   quantity            |             +------------------+\n"
        "                                 |   invoice_date        |\n"
        "                                 |   unit_price          |\n"
        "                                 |   total_amount        |\n"
        "                                 +-----------------------+"
    )
    add_code_block(er_ascii)

    # --- CHAPTER 4 ---
    add_chapter_heading("4", "Implementation")
    add_section_heading("4.1 TOOLS USED")
    doc.add_paragraph("Implemented in Python 3 using SQLite for database constraints, Pandas/NumPy for matrix operations, Scikit-Learn for K-Means clustering, and Streamlit/Tkinter for visual analytics.")

    add_section_heading("4.2 METHODOLOGY")
    doc.add_paragraph(
        "1. Schema & Trigger Compilation: Created Fact and Dimension tables in SQLite. An AFTER INSERT trigger automatically calculates total_amount = quantity * unit_price.\n"
        "2. Data Sanitization: Filtered missing Customer IDs and non-positive quantities/unit prices.\n"
        "3. RFM Formulation: Aggregated Recency (days since last purchase), Frequency (distinct invoice count), and Monetary (total spend).\n"
        "4. StandardScaler Normalization: Transformed features using z = (x - mean) / std to ensure equal distance weighting.\n"
        "5. K-Means Clustering: Grouped customers into 3 clusters, validated using Silhouette scoring, and assigned business segment labels."
    )

    add_section_heading("4.3 FUNCTIONAL MODULES")
    for mod in [
        "Module 1 (SQL Database Engine - run_sql_pipeline.py): Creates relational schema, triggers, and executes Fact-Dim joins, CLV aggregations, and subquery filters.",
        "Module 2 (ETL & ML Training Engine - 02_etl_rfm_ml.py): Cleans transaction records, engineers RFM metrics, fits K-Means, computes Silhouette score, and saves CSV outputs.",
        "Module 3 (Master Batch Pipeline - run_project.py): Complete automated execution logging KPIs to terminal and saving visual chart image outputs/dashboard_analytics.png.",
        "Module 4 (Streamlit Web Dashboard - dashboard/app.py): Real-time interactive browser dashboard with KPI metric cards and charts.",
        "Module 5 (Desktop GUI Interface - dashboard/app_gui.py): Standalone Tkinter desktop window with embedded Matplotlib figures."
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(mod)

    add_section_heading("4.4 TECHNICAL CODE ARTIFACTS")
    doc.add_paragraph("Relational Schema Definition & Automated Trigger (SQL):")
    sql_code = (
        "CREATE TABLE dim_customers (\n"
        "    customer_id INTEGER PRIMARY KEY,\n"
        "    country TEXT,\n"
        "    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ");\n\n"
        "CREATE TABLE dim_products (\n"
        "    stock_code TEXT PRIMARY KEY,\n"
        "    description TEXT,\n"
        "    unit_price REAL CHECK (unit_price >= 0)\n"
        ");\n\n"
        "CREATE TABLE fact_transactions (\n"
        "    invoice_no TEXT,\n"
        "    stock_code TEXT,\n"
        "    customer_id INTEGER,\n"
        "    quantity INTEGER,\n"
        "    invoice_date TEXT,\n"
        "    unit_price REAL,\n"
        "    total_amount REAL,\n"
        "    FOREIGN KEY(customer_id) REFERENCES dim_customers(customer_id),\n"
        "    FOREIGN KEY(stock_code) REFERENCES dim_products(stock_code)\n"
        ");\n\n"
        "CREATE TRIGGER trg_calc_transaction_total\n"
        "AFTER INSERT ON fact_transactions\n"
        "FOR EACH ROW\n"
        "BEGIN\n"
        "    UPDATE fact_transactions\n"
        "    SET total_amount = NEW.quantity * NEW.unit_price\n"
        "    WHERE rowid = NEW.rowid;\n"
        "END;"
    )
    add_code_block(sql_code)

    doc.add_paragraph("RFM Feature Engineering & K-Means Clustering (Python):")
    ml_code = (
        "# RFM Aggregation\n"
        "snapshot_date = df['InvoiceDate'].max() + pd.Timedelta(days=1)\n"
        "rfm = df.groupby('CustomerID').agg({\n"
        "    'InvoiceDate': lambda x: (snapshot_date - x.max()).days,\n"
        "    'InvoiceNo': 'nunique',\n"
        "    'TotalAmount': 'sum'\n"
        "}).reset_index()\n"
        "rfm.columns = ['CustomerID', 'Recency', 'Frequency', 'Monetary']\n\n"
        "# Normalization & Clustering\n"
        "scaler = StandardScaler()\n"
        "rfm_scaled = scaler.fit_transform(rfm[['Recency', 'Frequency', 'Monetary']])\n"
        "kmeans = KMeans(n_clusters=3, random_state=42, n_init=10)\n"
        "rfm['Cluster'] = kmeans.fit_predict(rfm_scaled)\n\n"
        "# Business Labeling\n"
        "summary = rfm.groupby('Cluster')['Monetary'].mean().sort_values()\n"
        "cluster_map = {\n"
        "    summary.index[0]: 'Low Value / Hibernating',\n"
        "    summary.index[1]: 'Mid-Tier / Potential Loyal',\n"
        "    summary.index[2]: 'High Value / Champions'\n"
        "}\n"
        "rfm['Segment'] = rfm['Cluster'].map(cluster_map)\n"
        "sil_score = silhouette_score(rfm_scaled, rfm['Cluster'])"
    )
    add_code_block(ml_code)

    add_section_heading("4.5 EXECUTION OUTPUTS & TERMINAL VERIFICATION")
    doc.add_paragraph("Master Pipeline Terminal Execution Log:")
    exec_log = (
        "=================================================================\n"
        "  ENTERPRISE RETAIL INTELLIGENCE (REAL ONLINE RETAIL DATASET)\n"
        "=================================================================\n\n"
        "[STEP 1] Loading pre-processed data from 'data/retail_clean.csv'...\n\n"
        "[STEP 2] Computing RFM Metrics & Training K-Means Cluster Model...\n"
        " -> Model Evaluation Silhouette Score: 0.3789\n"
        " -> Segment Summary:\n"
        "Mid-Tier / Potential Loyal    1412\n"
        "Low Value / Hibernating       1277\n"
        "High Value / Champions         641\n\n"
        "=================================================================\n"
        "                    BUSINESS METRICS SUMMARY\n"
        "=================================================================\n"
        " Total Revenue        : $121,314.76\n"
        " Total Invoices       : 2,733\n"
        " Unique Customers     : 3,330\n"
        " Average Order Value  : $24.26\n"
        "=================================================================\n"
        "[STEP 3] Generating Visual Plots... [outputs/dashboard_analytics.png SAVED]"
    )
    add_code_block(exec_log)

    # --- CHAPTER 5 ---
    add_chapter_heading("5", "Results and Discussion")
    add_section_heading("5.1 OUTPUT / BUSINESS KPIS & CLUSTER COHORTS")
    doc.add_paragraph("The end-to-end pipeline was executed on transaction datasets, producing verified metrics across business operations and machine learning clustering:")

    table2 = doc.add_table(rows=4, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table2)
    h2 = ["Customer Segment", "Count", "Percentage", "Strategic Business Profile"]
    for i, h in enumerate(h2):
        cell = table2.cell(0, i)
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    seg_data = [
        ("Mid-Tier / Potential Loyal", "1,412", "42.4%", "Moderate recency and order frequency; ideal for product cross-selling."),
        ("Low Value / Hibernating", "1,277", "38.3%", "High recency (inactive) with low spend; targets for email re-engagement."),
        ("High Value / Champions", "641", "19.3%", "Highest spend and frequent orders; prime candidates for VIP loyalty perks.")
    ]
    for row_idx, data in enumerate(seg_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table2.cell(row_idx, col_idx)
            set_cell_background(cell, "FFFFFF" if row_idx % 2 != 0 else "F8F9F9")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
    doc.add_paragraph()

    add_section_heading("5.2 CHALLENGES FACED")
    doc.add_paragraph(
        "1. Feature Scale Dominance: Monetary spend ranged in thousands while Recency ranged in days. Without scaling, Euclidean distance was distorted by Monetary values. Resolved using StandardScaler.\n"
        "2. File Locks on Windows OS: Concurrent file reads/writes in Excel caused PermissionError. Resolved using fallback paths and in-memory execution.\n"
        "3. UI Latency: Model re-training on dashboard reload caused UI lag. Decoupled offline batch training from presentation layers for instant load times."
    )

    add_section_heading("5.3 LEARNINGS & TECHNICAL INSIGHTS")
    doc.add_paragraph(
        "Gained hands-on experience designing Star Schema relational architectures with automated triggers, extracting RFM metrics, "
        "training distance-based machine learning models, and building multi-platform analytical user interfaces."
    )

    # --- CHAPTER 6 ---
    add_chapter_heading("6", "Conclusion")
    add_section_heading("6.1 SUMMARY")
    doc.add_paragraph(
        "The Enterprise Retail Intelligence & Customer Segmentation System successfully bridges relational database engineering, "
        "RFM feature transformation, unsupervised K-Means clustering (Silhouette = 0.3789), and modern multi-interface visualization. "
        "All project objectives were completed and verified."
    )

    add_section_heading("6.2 FUTURE ENHANCEMENTS")
    doc.add_paragraph(
        "1. Dynamic Hyperparameter Sweeps: Automate optimal K selection using dynamic Elbow and Silhouette sweeps (K=2 to 10).\n"
        "2. Cloud Database Migration: Scale storage to cloud data warehouses like Snowflake or PostgreSQL.\n"
        "3. Churn Prediction Engine: Implement supervised classification (XGBoost / Random Forest) to predict churn probabilities."
    )

    # --- ANNEXURES ---
    doc.add_page_break()
    p_anxa = doc.add_paragraph()
    r_anxa = p_anxa.add_run("ANNEXURE A: MASTER BATCH PIPELINE SOURCE CODE")
    r_anxa.bold = True
    r_anxa.font.size = Pt(13)
    r_anxa.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph("Standalone Python Master Pipeline (run_project.py):")
    full_code = (
        "import os\nimport pandas as pd\nimport numpy as np\nimport matplotlib.pyplot as plt\n"
        "from sklearn.preprocessing import StandardScaler\nfrom sklearn.cluster import KMeans\nfrom sklearn.metrics import silhouette_score\n\n"
        "print('='*65)\nprint('  ENTERPRISE RETAIL INTELLIGENCE (REAL ONLINE RETAIL DATASET)')\nprint('='*65)\n\n"
        "csv_clean_path = 'data/retail_clean.csv'\n"
        "df = pd.read_csv(csv_clean_path)\n"
        "df['InvoiceDate'] = pd.to_datetime(df['InvoiceDate'])\n\n"
        "# RFM Aggregation\n"
        "snapshot_date = df['InvoiceDate'].max() + pd.Timedelta(days=1)\n"
        "rfm = df.groupby('CustomerID').agg({\n"
        "    'InvoiceDate': lambda x: (snapshot_date - x.max()).days,\n"
        "    'InvoiceNo': 'nunique',\n"
        "    'TotalAmount': 'sum'\n"
        "}).reset_index()\n"
        "rfm.columns = ['CustomerID', 'Recency', 'Frequency', 'Monetary']\n\n"
        "# Standard Scaling & K-Means\n"
        "scaler = StandardScaler()\n"
        "rfm_scaled = scaler.fit_transform(rfm[['Recency', 'Frequency', 'Monetary']])\n"
        "kmeans = KMeans(n_clusters=3, random_state=42, n_init=10)\n"
        "rfm['Cluster'] = kmeans.fit_predict(rfm_scaled)\n\n"
        "summary = rfm.groupby('Cluster')['Monetary'].mean().sort_values()\n"
        "cluster_map = {summary.index[0]: 'Low Value / Hibernating', summary.index[1]: 'Mid-Tier / Potential Loyal', summary.index[2]: 'High Value / Champions'}\n"
        "rfm['Segment'] = rfm['Cluster'].map(cluster_map)\n"
        "sil_score = silhouette_score(rfm_scaled, rfm['Cluster'])\n\n"
        "os.makedirs('outputs', exist_ok=True)\n"
        "rfm.to_csv('data/customer_segments.csv', index=False)\n"
        "print(f' -> Model Evaluation Silhouette Score: {sil_score:.4f}')\n"
        "print(f' Total Revenue: ${df[\"TotalAmount\"].sum():,.2f}')\n"
    )
    add_code_block(full_code)

    doc.add_page_break()
    p_anxb = doc.add_paragraph()
    r_anxb = p_anxb.add_run("ANNEXURE B: PRESENTATION DECK STRUCTURE")
    r_anxb.bold = True
    r_anxb.font.size = Pt(13)
    r_anxb.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    ppt_struct = (
        "SLIDE 1: Title Slide (Project Title, Ayush Raj, Reg No: 12315673, CSE343/CSE443, LPU)\n"
        "SLIDE 2: Problem Statement & Motivation (High-volume retail noise, customer attrition, need for segmentation)\n"
        "SLIDE 3: System Architecture & Star Schema (Fact/Dim tables, trg_calc_transaction_total automated trigger)\n"
        "SLIDE 4: RFM Feature Engineering & K-Means Pipeline (StandardScaler z-scores, K=3 clustering, Silhouette: 0.3789)\n"
        "SLIDE 5: Business KPIs & Experimental Results (Revenue: $121,314.76, 2,733 Invoices, 3,330 Customers, AOV: $24.26)\n"
        "SLIDE 6: Multi-Tier Interfaces (Streamlit Web Dashboard, Tkinter Desktop GUI, Headless Runner)\n"
        "SLIDE 7: Conclusion & Future Scope (Elbow sweeps, Cloud Snowflake migration, Churn prediction)"
    )
    add_code_block(ppt_struct)

    doc.save(OUTPUT_PATH)
    print(f"\n[SUCCESS] Formatted Word Document generated at:")
    print(f" -> '{OUTPUT_PATH}'")

if __name__ == "__main__":
    build_professional_report()